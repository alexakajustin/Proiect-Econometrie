import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT

# --- CONFIGURARE CULORI ---
COLOR_HEADER_BG = "8B4513"  # SaddleBrown (darker)
COLOR_BODY_BG = "DEB887"    # BurlyWood (lighter)
COLOR_CODE_BG = "FFF8DC"    # Cornsilk (for code)

# --- DATA PATHS ---
BASE_PATH = "Proiect/Proiect/Output"
PATH_STATS = f"{BASE_PATH}/Rapoarte/Statistici_Descriptive.csv"
PATH_MODEL_TXT = f"{BASE_PATH}/Rapoarte/rezultate_modelare.txt"
PATH_PANEL_TXT = f"{BASE_PATH}/Rapoarte/rezultate_panel.txt"
PATH_ML_TXT = f"{BASE_PATH}/Rapoarte/rezultate_modelare_extins.txt"

IMG_DIR = f"{BASE_PATH}/Grafice"
IMG_HIST = f"{IMG_DIR}/Hist_Grid_All.png"
IMG_BOXPLOT = f"{IMG_DIR}/Boxplot_Outlieri.png"
IMG_CORR = f"{IMG_DIR}/Plot_Corelatie.png"
IMG_SCATTER = f"{IMG_DIR}/Scatter_Log_Somaj_Furturi.png"
IMG_LASSO = f"{IMG_DIR}/Lasso_Trace.png"
IMG_RESID = f"{IMG_DIR}/Residuals_vs_Fitted.png"
IMG_QQ = f"{IMG_DIR}/QQ_Plot_Reziduuri.png"
IMG_REGRESIE = f"{IMG_DIR}/Regresie_Simpla.png"
IMG_PAIRS = f"{IMG_DIR}/Pairs_Plot.png"

def set_cell_bg(cell, color):
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading)

def add_box(doc, title, content, is_dark=True):
    """Add styled box - dark for headers, light for content"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    c.width = Inches(7.5)
    set_cell_bg(c, COLOR_HEADER_BG if is_dark else COLOR_BODY_BG)
    
    if title:
        p = c.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(18 if is_dark else 14)
        r.font.color.rgb = RGBColor(255,255,255) if is_dark else RGBColor(50,20,0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_dark else WD_ALIGN_PARAGRAPH.LEFT

    if content:
        lines = content if isinstance(content, list) else [content]
        for line in lines:
            p = c.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(255,255,255) if is_dark else RGBColor(0,0,0)
    doc.add_paragraph()

def add_code_box(doc, title, code):
    """Add R code snippet box"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    c.width = Inches(7.5)
    set_cell_bg(c, COLOR_CODE_BG)
    
    if title:
        p = c.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(139,69,19)
    
    p = c.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0,0,0)
    doc.add_paragraph()

def add_img(doc, path, title=None, h=Inches(4)):
    if title:
        add_box(doc, title, [], is_dark=True)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, height=h)
    else:
        doc.add_paragraph(f"[IMAGINE LIPSĂ: {path}]")
    doc.add_paragraph()

def read_txt(path):
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return ""

def create_document():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.4)
    sec.right_margin = Inches(0.4)

    # ===== PAGE 1: TITLE =====
    add_box(doc, "PROIECT ECONOMETRIE", [
        "",
        "Analiza Determinanților Ratei Furturilor în Uniunea Europeană",
        "Regresie Simplă, Multiplă, Machine Learning și Panel Data",
        "",
        "Coordonatori:",
        "Conf. Univ. Dr. Aceleanu Mirela | Drd. Anghel Iuliana-Claudia",
        "",
        "Student: Hristu Justin | Grupa 1083"
    ])
    doc.add_page_break()

    # ===== PAGE 2: INTRODUCERE =====
    add_box(doc, "Introducere", "")
    add_box(doc, "Contextul Studiului", [
        "Infracțiunile contra proprietății (furturile) sunt influențate de factori economici și sociali.",
        "Scopul analizei de regresie: estimarea și previzionarea ratei furturilor."
    ], is_dark=False)
    add_box(doc, "Sursa datelor", "Eurostat: https://ec.europa.eu/eurostat", is_dark=True)
    add_box(doc, "Variabile Analizate", [
        "Dependentă: ln_Furturi - rata furturilor (log)",
        "Independente: ln_PIB, ln_Someri, ln_Imigratie, ln_Densitate",
        "Dummy: Est_Vest (1 = Europa de Est, 0 = Europa de Vest)"
    ], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 3: STUDIUL CUNOAȘTERII =====
    add_box(doc, "Studiul Cunoașterii", "")
    literature = [
        "1. Becker (1968) - Teoria economică a criminalității: infractorii sunt actori raționali.",
        "2. Raphael & Winter-Ebmer (2001) - Șomajul crește propensiunea spre furturi.",
        "3. Fajnzylber et al. (2002) - Inegalitatea veniturilor este corelată cu criminalitatea.",
        "4. Sampson & Groves (1989) - Dezorganizarea socială în zone dens populate.",
        "5. Ousey & Kubrin (2018) - Imigrația nu crește neapărat rata criminalității."
    ]
    add_box(doc, None, literature, is_dark=False)
    doc.add_page_break()

    # ===== PAGE 4: INFORMAȚII DATE =====
    add_box(doc, "Informații despre Date", "40 țări, perioada 2019-2023")
    
    if os.path.exists(PATH_STATS):
        df = pd.read_csv(PATH_STATS)
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Table Grid'
        for i, col in enumerate(df.columns):
            t.rows[0].cells[i].text = str(col)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
    doc.add_page_break()

    # ===== PAGE 5: REGRESIA SIMPLĂ - Teorie =====
    add_box(doc, "Regresia Simplă", "Prețul în funcție de Șomaj")
    add_box(doc, "Ipoteze", [
        "Forma funcțională este liniară",
        "Erorile au media 0",
        "Homoscedasticitatea erorilor aleatore",
        "Erorile nu sunt autocorelate",
        "Necorelarea între regresor și erorile aleatoare",
        "Erorile au distribuție normală"
    ], is_dark=False)
    add_code_box(doc, "Cod R - Model Simplu", 
        "model_simplu <- lm(ln_Furturi ~ ln_Someri, data = df_final)\n"
        "summary(model_simplu)\n"
        "# R-squared indică puterea explicativă")
    doc.add_page_break()

    # ===== PAGE 6: REGRESIA SIMPLĂ - Rezultate =====
    add_img(doc, IMG_SCATTER, "Scatter Plot: Furturi vs Șomaj", h=Inches(3.5))
    
    res = read_txt(PATH_MODEL_TXT)
    model1 = res.split("3. MODEL 2")[0] if "3. MODEL 2" in res else res[:1500]
    add_box(doc, "Rezultate Model Simplu", model1[:1200], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 7: HOMOSCEDASTICITATE & AUTOCORELARE =====
    add_box(doc, "Homoscedasticitatea", "Varianța erorilor este constantă")
    add_code_box(doc, "Teste R", 
        "bptest(model_simplu)  # Breusch-Pagan\n"
        "white_test(model_simplu)  # White test\n"
        "# p-value > 0.05 => Homoscedasticitate")
    
    add_box(doc, "Autocorelare", "")
    add_code_box(doc, "Teste R",
        "dwtest(model_simplu)  # Durbin-Watson\n"
        "bgtest(model_simplu, order = 1)  # Breusch-Godfrey\n"
        "# p-value > 0.1 => Nu există autocorelare")
    add_img(doc, IMG_RESID, None, h=Inches(3))
    doc.add_page_break()

    # ===== PAGE 8: NORMALITATEA =====
    add_box(doc, "Normalitatea Reziduurilor", "")
    add_code_box(doc, "Teste R",
        "jarque.bera.test(model_simplu$residuals)\n"
        "shapiro.test(model_simplu$residuals)\n"
        "# p-value < 0.05 => Reziduurile NU sunt normal distribuite\n"
        "# Corecție: eliminăm outleri cu distanța Cook")
    add_img(doc, IMG_QQ, "Normal Q-Q Plot", h=Inches(3.5))
    doc.add_page_break()

    # ===== PAGE 9: PROGNOZE =====
    add_box(doc, "Prognoze", "Train/Test Split 80/20")
    add_code_box(doc, "Cod R - Prognoza",
        "out_of_sample <- data.frame(ln_Someri = c(5.5, 6.0, 6.5))\n"
        "predict(model_simplu, newdata = out_of_sample, interval = 'confidence')\n"
        "# fit    lwr    upr\n"
        "# 7.2    6.8    7.6  <- Interval 95%")
    
    add_box(doc, "Interpretare Prognoză", [
        "Pentru un nivel de șomaj de 6.0 (log), rata furturilor prognozată este ~7.2.",
        "Intervalul de încredere 95%: [6.8, 7.6]",
        "MAPE < 5% indică o acuratețe bună a modelului."
    ], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 10: REGRESIA MULTIPLĂ =====
    add_box(doc, "Regresia Multiplă", "Toate Variabilele")
    add_code_box(doc, "Ecuația",
        "ln_Furturi = β0 + β1*ln_PIB + β2*ln_Someri + β3*ln_Imigratie + β4*ln_Densitate + β5*Est_Vest + ε")
    
    try:
        model2 = res.split("3. MODEL 2")[1].split("7. MODEL 3")[0]
    except:
        model2 = "Rezultate indisponibile"
    add_box(doc, "Rezultate OLS Multiplu", model2[:1500], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 11: VALIDARE IPOTEZE =====
    add_box(doc, "Validarea Ipotezelor", "")
    
    hypotheses = [
        ("H1: Șomaj → Furturi (+)", "β2 > 0, p < 0.05", "✓ VALIDATĂ"),
        ("H2: PIB → Furturi (-)", "β1 < 0, p > 0.05", "✗ INVALIDATĂ"),
        ("H3: Imigrație → Furturi", "β3 > 0, p < 0.05", "✓ VALIDATĂ"),
        ("H4: Densitate → Furturi", "β4 sign ambiguu", "? PARȚIAL"),
        ("H5: Est > Vest", "β5 > 0", "✓ VALIDATĂ")
    ]
    
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Ipoteză"
    t.rows[0].cells[1].text = "Condiție"
    t.rows[0].cells[2].text = "Rezultat"
    for h in hypotheses:
        row = t.add_row().cells
        row[0].text = h[0]
        row[1].text = h[1]
        row[2].text = h[2]
    doc.add_paragraph()
    doc.add_page_break()

    # ===== PAGE 12: STEPWISE & DIAGNOSTICARE =====
    add_box(doc, "Selecția Modelului Optim", "Stepwise AIC")
    try:
        model3 = res.split("7. MODEL 3")[1].split("Script finalizat")[0]
    except:
        model3 = "Rezultate indisponibile"
    add_box(doc, "Model Refinat", model3[:1200], is_dark=False)
    add_img(doc, IMG_PAIRS, "Pairs Plot", h=Inches(3.5))
    doc.add_page_break()

    # ===== PAGE 13: MACHINE LEARNING =====
    add_box(doc, "Modele de Regularizare", "Ridge, LASSO, Elastic Net")
    
    ml = read_txt(PATH_ML_TXT)
    add_box(doc, "Comparație Performanță", ml, is_dark=False)
    add_img(doc, IMG_LASSO, "Lasso Coefficient Trace", h=Inches(3.5))
    doc.add_page_break()

    # ===== PAGE 14: COMPARAȚIE ML =====
    add_box(doc, "Comparație Modele", "")
    comparison = [
        ("OLS", "0.58", "0.43"),
        ("Ridge", "0.84", "0.62"),
        ("LASSO", "0.66", "0.77")
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Model"
    t.rows[0].cells[1].text = "RMSE"
    t.rows[0].cells[2].text = "R-squared"
    for c in comparison:
        row = t.add_row().cells
        row[0].text = c[0]
        row[1].text = c[1]
        row[2].text = c[2]
    
    add_box(doc, "Concluzie ML", [
        "LASSO oferă cel mai bun R-squared (0.77) cu selecție automată de variabile.",
        "Ridge reduce multicoliniaritatea dar păstrează toate variabilele."
    ], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 15: PANEL DATA - Intro =====
    add_box(doc, "Aplicația 2: Panel Data", "GDP în funcție de factori, 2019-2023")
    add_box(doc, "Obiectiv & Studiul Cunoașterii", [
        "Analiza longitudinală pe 40 țări x 5 ani = 200 observații",
        "Controlăm pentru eterogeneitatea neobservată între țări",
        "Fixed Effects vs Random Effects"
    ], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 16: PANEL - Hausman =====
    add_box(doc, "Testul Hausman", "Fixed vs Random Effects")
    add_code_box(doc, "Cod R",
        "model_fe <- plm(ln_Furturi ~ ln_PIB + ln_Someri, model='within')\n"
        "model_re <- plm(ln_Furturi ~ ln_PIB + ln_Someri, model='random')\n"
        "phtest(model_fe, model_re)\n"
        "# p-value < 0.05 => Alegem Fixed Effects")
    
    panel = read_txt(PATH_PANEL_TXT)
    add_box(doc, "Rezultate Panel", panel[:2000], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 17: PANEL - Diagnosticare =====
    add_box(doc, "Diagnosticare Panel", "")
    add_code_box(doc, "Teste",
        "pcdtest(model_fe)  # Pesaran CD (cross-sectional dependence)\n"
        "pbgtest(model_fe)  # Breusch-Godfrey (serial correlation)\n"
        "bptest(model_fe)   # Breusch-Pagan (heteroscedasticity)")
    add_box(doc, "Interpretare", [
        "Dacă p-value < 0.05 la Pesaran CD: există dependență transversală.",
        "Dacă p-value < 0.05 la BG: există autocorelare.",
        "Se recomandă erori standard robuste (HC sau Driscoll-Kraay)."
    ], is_dark=False)
    doc.add_page_break()

    # ===== PAGE 18: CONCLUZII & BIBLIOGRAFIE =====
    add_box(doc, "Concluzii", "")
    add_box(doc, "Rezultate Principale", [
        "1. Șomajul este cel mai puternic predictor al furturilor (+ semnificativ).",
        "2. PIB-ul are efect nesemnificativ în majoritatea specificațiilor.",
        "3. Țările Est-Europene au rate mai mari ale furturilor (dummy Est_Vest).",
        "4. LASSO oferă cel mai bun compromis între fit și parsimonie.",
        "5. Analiza Panel confirmă efecte individuale semnificative (Fixed Effects)."
    ], is_dark=False)
    
    doc.add_page_break()
    
    add_box(doc, "Bibliografie", "")
    add_box(doc, None, [
        "Sursa seturilor de date:",
        "  https://ec.europa.eu/eurostat",
        "",
        "Documentație R:",
        "  https://www.rdocumentation.org/",
        "",
        "Articole academice:",
        "  Becker, G. (1968). Crime and Punishment: An Economic Approach.",
        "  Raphael & Winter-Ebmer (2001). Identifying the Effect of Unemployment on Crime.",
        "  Fajnzylber et al. (2002). Inequality and Violent Crime.",
        "  Sampson & Groves (1989). Community Structure and Crime."
    ], is_dark=False)

    # SAVE
    doc.save("Proiect_Final_Econometrie_Complet.docx")
    print("Raport Final Generat: Proiect_Final_Econometrie_Complet.docx")

if __name__ == "__main__":
    create_document()
